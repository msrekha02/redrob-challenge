"""
jd_processor.py — Dynamic JD Profile + 4-Way Multi-Cast (v4)


"""

from __future__ import annotations

import logging
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import numpy as np

sys.path.insert(0, str(Path(__file__).parent))
import config as cfg

log = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
# JDProfile
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class JDProfile:
    """Auto-computed descriptor of any JD. Passed through the full pipeline."""
    # Detected role properties (reported for logging; weights below are a blend)
    role_category:        str            # highest-similarity category, for display only
    category_confidence:  float          # that category's raw similarity, for display only
    category_low_confidence: bool        # True if no category resembles the JD well
    seniority_level:      float
    technical_depth:      float
    urgency:              float

    # Dynamic weights — similarity-weighted BLEND across all categories,
    # not a single bucket's fixed profile
    role_weight:               float
    cap_weight:                float
    avail_importance:          float
    loc_importance:            float
    domain_mismatch_threshold: float      # auto-expands under low confidence — see fix #3

    # Continuous-penalty gates (renamed from *_dq_applies — see module docstring)
    consulting_penalty_active: bool
    research_penalty_active:   bool
    production_is_required:    bool

    preferred_cities: set[str] = field(default_factory=set)
    yoe_min: int | None = None
    yoe_max: int | None = None            # open-ended -> YOE_OPEN_ENDED_CEILING (99), not None
    nlp_ir_required: bool = False

    # Text-native notice-period extraction (None if JD doesn't state one anywhere)
    notice_period_days_stated: int | None = None

    # BLENDED semantic anchor vector — weighted average of all 10 category
    # anchor embeddings, weighted by the same softmax used for the numeric
    # blend. Used for the domain-mismatch check in guardrails.py.
    category_anchor_vec: np.ndarray = field(default_factory=lambda: np.zeros(cfg.EMBED_DIM, dtype=np.float32))

    flags: dict[str, bool] = field(default_factory=dict)


# ─────────────────────────────────────────────────────────────────────────────
# Detection helpers — pure functions
# ─────────────────────────────────────────────────────────────────────────────

_SENIORITY_EXEC_KW   = frozenset(["vp of", "vice president", "director of", "head of", "chief ", "c-suite"])
_SENIORITY_SENIOR_KW = frozenset(["principal", "staff engineer", "founding team", "founding member", "senior", " sr ", "sr."])
_SENIORITY_MID_KW    = frozenset(["mid-level", "intermediate", "3-5 years", "3 to 5", "associate"])
_SENIORITY_JUNIOR_KW = frozenset(["junior", "jr.", "entry level", "entry-level", "0-2 years", "fresher", "new graduate"])

def _detect_seniority(text_lower: str) -> float:
    if any(kw in text_lower for kw in _SENIORITY_EXEC_KW):   return 1.00
    if any(kw in text_lower for kw in _SENIORITY_SENIOR_KW): return 0.78
    if any(kw in text_lower for kw in _SENIORITY_MID_KW):    return 0.50
    if any(kw in text_lower for kw in _SENIORITY_JUNIOR_KW): return 0.22
    return 0.65


_TECH_DEEP_SIGNALS = frozenset([
    "algorithm", "latency", "throughput", "distributed system", "machine learning",
    "neural network", "embeddings", "vector", "kubernetes", "microservice",
    "database", "api", "backend", "compiler", "runtime", "system design",
    "architecture", "pipeline", "inference", "model", "python", "java", "c++",
    "sql", "nosql", "ci/cd", "devops", "framework", "tensor", "gradient",
    "fine-tuning", "tokenizer", "retrieval", "ranking",
])
_SOFT_SIGNALS = frozenset([
    "stakeholder", "strategy", "roadmap", "campaign", "brand", "customer",
    "sales", "marketing", "growth", "revenue", "fundraising", "partnership",
    "communication", "negotiation", "business development", "influencing",
])

def _detect_technical_depth(text_lower: str) -> float:
    tech = sum(1 for s in _TECH_DEEP_SIGNALS if s in text_lower)
    soft = sum(1 for s in _SOFT_SIGNALS if s in text_lower)
    return (tech / (tech + soft)) if (tech + soft) > 0 else 0.50


_URGENCY_HIGH = frozenset(["immediate joiner", "immediate start", "asap", "urgently"])
_URGENCY_MED  = frozenset(["sub-30", "30 days notice", "buy out notice", "notice period < 30", "sub 30"])
_URGENCY_LOW  = frozenset(["flexible start", "notice period negotiable", "3 months notice", "6 months"])

def _detect_urgency(text_lower: str) -> float:
    if any(kw in text_lower for kw in _URGENCY_HIGH): return 1.00
    if any(kw in text_lower for kw in _URGENCY_MED):  return 0.65
    if any(kw in text_lower for kw in _URGENCY_LOW):  return 0.20
    return 0.40


def _extract_dq_conditions(text_lower: str) -> dict[str, bool]:
    """
    Extract explicit penalty-applicability conditions from JD text.
    Renamed semantics: these no longer mean "hard-disqualify" — they mean
    "the continuous recency-weighted penalty for this signal should apply
    at all for this JD." See guardrails.py for the actual scoring.
    """
    dq_verbs = r"(?:will\s+not|won.t|do\s+not|don.t|cannot|not\s+a\s+fit|disqualif|reject|eliminat)"

    consulting_relevant = bool(re.search(
        rf"(?:consulting|services?\s+firm|tcs|infosys|wipro|accenture|cognizant).{{0,200}}{dq_verbs}|"
        rf"{dq_verbs}.{{0,200}}(?:consulting|services?\s+firm|only\s+work)",
        text_lower, re.DOTALL
    )) or bool(re.search(
        r"people\s+who\s+have\s+only\s+work\w+\s+at\s+consulting",
        text_lower,
    ))

    research_relevant = bool(re.search(
        rf"(?:pure\s+research|academic\s+lab|research.only|without.*?production).{{0,200}}{dq_verbs}|"
        rf"{dq_verbs}.{{0,200}}(?:pure\s+research|academic|no\s+production)",
        text_lower, re.DOTALL
    ))

    production_required = bool(re.search(
        r"\b(?:production|deployed|real\s+users|at\s+scale)\b", text_lower
    ))

    return {
        "consulting_penalty_active": consulting_relevant,
        "research_penalty_active":   research_relevant,
        "production_is_required":    production_required,
    }


def _extract_preferred_cities(jd_text: str) -> set[str]:
    text_lower = jd_text.lower()
    found: set[str] = set()
    for city in cfg.KNOWN_CITY_POOL:
        if city not in text_lower:
            continue
        idx     = text_lower.find(city)
        context = text_lower[max(0, idx - 100): idx + 120]
        if any(pos in context for pos in cfg.CITY_POSITIVE_CONTEXTS):
            found.add(city)
    return found if found else set(cfg.PREFERRED_CITIES)


# ─────────────────────────────────────────────────────────────────────────────
# FIX #1 — YoE range: multi-match scan, widest bounds, numeric open-ended ceiling
# ─────────────────────────────────────────────────────────────────────────────

# Sentinel ceiling for an open-ended upper bound ("10+ years" -> 99, not
# None). Keeps every downstream numeric comparison and format call working
# without a None-guard for the upper bound specifically. yoe_min has no
# equivalent sentinel — an unstated FLOOR has no sensible non-None
# substitute the way an unstated CEILING does (99 reads naturally as
# "effectively unbounded"; a fake "0" floor would silently change meaning).
YOE_OPEN_ENDED_CEILING: int = 99

# Sanity bound for any single extracted year value — rejects clearly
# nonsensical matches (typos, unrelated large numbers) without crashing.
_YOE_SANITY_MAX: int = 60


# Context phrases that, when found near a matched "X years" span, mean the
# match is NOT about total years of experience — job-tenure commitments
# ("plans to be here for 3+ years") or topic-specific durations
# ("closed-source systems for 5+ years") use the same "N+ years" surface
# form as a genuine experience floor but mean something else entirely.
# Discovered by testing against the full real JD text, not a short
# synthetic snippet — short snippets don't surface this kind of incidental
# collision with other "years"-shaped phrases elsewhere in a long JD.
_YOE_EXCLUSION_CONTEXT: tuple[str, ...] = (
    "plan to be here", "plans to be here", "to stay", "commit",
    "tenure", "closed-source", "closed source", "proprietary system",
)


def _yoe_context_is_valid(text_lower: str, start: int, end: int) -> bool:
    """
    Rejects a matched 'X years' span if it sits in the SAME SENTENCE as
    vocabulary describing something other than total years of experience
    (see _YOE_EXCLUSION_CONTEXT). Sentence-bounded rather than a fixed
    character window: two genuinely separate mentions in different
    sentences shouldn't cross-contaminate each other just because they
    happen to sit within a few dozen characters of one another (which can
    legitimately happen in a JD with short, punchy sentences).
    """
    boundary_chars = (".", "!", "?", "\n")
    sentence_start = max(
        (text_lower.rfind(c, 0, start) for c in boundary_chars), default=-1
    ) + 1
    end_positions = [text_lower.find(c, end) for c in boundary_chars]
    valid_ends = [e for e in end_positions if e != -1]
    sentence_end = min(valid_ends) if valid_ends else len(text_lower)
    sentence = text_lower[sentence_start:sentence_end]
    return not any(ctx in sentence for ctx in _YOE_EXCLUSION_CONTEXT)


def _detect_yoe_range(jd_text: str) -> tuple[int | None, int | None]:
    """
    Scans the FULL JD text for every YoE-range-like mention — not just the
    first match — across three patterns (explicit ranges, open-ended
    minimums, open-ended maximums), then aggregates to the WIDEST bounds:
    min(all minimums found), max(all maximums found). A JD that states
    "5-9 years" in the main description and "3+ years in distributed
    systems" in a sub-bullet ends up as (3, 9), not arbitrarily narrowed to
    whichever phrase the old single re.search happened to hit first.
    Erring wide here is deliberate: a narrower-than-intended range would
    soft-penalize candidates the JD never meant to exclude.
    """
    text_lower = jd_text.lower()
    mins: list[int] = []
    maxs: list[int] = []

    # Tier A: explicit "X-Y years" / "X to Y years" ranges — ALL matches.
    # "of experience/exp/total" suffix is optional (not required) so
    # phrasings like "5-9 years in NLP roles" are still captured — widening
    # recall is the explicit goal here, not narrowing it.
    for m in re.finditer(
        r"(\d{1,2})\s*[-–to]+\s*(\d{1,2})\s*(?:years?|yrs?)(?:\s+of)?\s*(?:experience|exp|total)?",
        text_lower,
    ):
        if not _yoe_context_is_valid(text_lower, m.start(), m.end()):
            continue
        lo, hi = int(m.group(1)), int(m.group(2))
        if 0 <= lo <= hi <= _YOE_SANITY_MAX:
            mins.append(lo)
            maxs.append(hi)

    # Tier B: open-ended minimums — "minimum X years", "at least X years", "X+ years"
    for m in re.finditer(
        r"(?:minimum|at\s+least|min\.?)\s*(\d{1,2})\s*(?:years?|yrs?)|(\d{1,2})\+\s*(?:years?|yrs?)",
        text_lower,
    ):
        if not _yoe_context_is_valid(text_lower, m.start(), m.end()):
            continue
        n = int(m.group(1) or m.group(2))
        if 0 <= n <= _YOE_SANITY_MAX:
            mins.append(n)
            # Deliberately does not append to maxs — this pattern states no
            # upper bound, which is exactly what makes it "open-ended".

    # Tier C: open-ended maximums — "up to X years", "no more than X years"
    for m in re.finditer(
        r"(?:up\s+to|no\s+more\s+than|maximum|max\.?)\s*(\d{1,2})\s*(?:years?|yrs?)",
        text_lower,
    ):
        if not _yoe_context_is_valid(text_lower, m.start(), m.end()):
            continue
        n = int(m.group(1))
        if 0 <= n <= _YOE_SANITY_MAX:
            maxs.append(n)

    if not mins and not maxs:
        return None, None

    yoe_min = min(mins) if mins else None
    # A real maximum was found somewhere -> use the widest one. Otherwise,
    # if we have a minimum but no maximum anywhere in the text, this is
    # genuinely open-ended ("10+ years") -> numeric ceiling, not None.
    yoe_max = max(maxs) if maxs else (YOE_OPEN_ENDED_CEILING if mins else None)

    return yoe_min, yoe_max


def _format_yoe_range(yoe_min: int | None, yoe_max: int | None) -> str:
    """
    Human-readable YoE range for logging/summary. Shows "X+" for the
    open-ended sentinel ceiling rather than a confusing "X–99".
    """
    min_str = _safe_fmt(yoe_min)
    if yoe_max == YOE_OPEN_ENDED_CEILING:
        return f"{min_str}+"
    return f"{min_str}–{_safe_fmt(yoe_max)}"


def _detect_nlp_ir_required(text_lower: str) -> bool:
    NLP_IR_REQUIRED_SIGNALS = frozenset([
        "nlp", "natural language processing", "information retrieval",
        "ranking system", "retrieval system", "search system",
        "recommendation system", "embedding", "semantic search",
    ])
    return any(sig in text_lower for sig in NLP_IR_REQUIRED_SIGNALS)


# ─────────────────────────────────────────────────────────────────────────────
# FIX #2 — Notice period: explicit contextual cascade
# ─────────────────────────────────────────────────────────────────────────────

_NOTICE_DAYS_MIN: int = 0     # exclusive lower bound
_NOTICE_DAYS_MAX: int = 365   # inclusive upper bound


def _bounded_days(value: int) -> int | None:
    """Sanity-bound any extracted day-count; returns None if out of range
    rather than silently accepting (or crashing on) a nonsensical match."""
    return value if _NOTICE_DAYS_MIN < value <= _NOTICE_DAYS_MAX else None


def _extract_notice_period_days(jd_text: str) -> int | None:
    """
    Contextual cascade, tried in priority order from most explicit to most
    general. Returns the first valid match found across all tiers; None
    only if every tier finds nothing in the entire JD text — a clean,
    deliberate "JD doesn't state a notice period" signal (the "defensive
    baseline"), not a crash or a guess. Every tier is independently bounded
    via _bounded_days, so this function can never raise and can never
    return an out-of-range number.

    Tier 1 — explicit compact tags: "sub-30", "sub-45-day", "under 30 days"
    Tier 2 — sentence-style explicit statements: "notice period of 45
              days", "notice period: 60 days", "30 day notice period"
    Tier 3 — colloquial timeframes, converted to days: "immediate
              joiner"/"asap"/"available immediately", "within N weeks",
              "N month(s) notice"
    Tier 4 — generic catch-all: any "N days" mention within a short window
              of notice/join/buyout context words (the original
              single-pattern approach — kept as the most permissive,
              lowest-priority tier, since a naive regex like a bare
              "notice period \\d+" misses colloquial unstructured phrasing
              on its own)
    """
    text_lower = jd_text.lower()

    # ── Tier 1: explicit compact tags ────────────────────────────────────────
    m = re.search(r"sub-?\s*(\d{1,3})\s*-?\s*days?\b", text_lower)
    if m:
        result = _bounded_days(int(m.group(1)))
        if result is not None:
            return result

    m = re.search(r"(?:less\s+than|under|<)\s*(\d{1,3})\s*days?\b", text_lower)
    if m:
        result = _bounded_days(int(m.group(1)))
        if result is not None:
            return result

    # ── Tier 2: sentence-style explicit statements ───────────────────────────
    m = re.search(
        r"notice\s+period\s*(?:of|is|:)?\s*(\d{1,3})\s*-?\s*days?\b|"
        r"(\d{1,3})\s*-?\s*days?\s+notice\s+period\b",
        text_lower,
    )
    if m:
        result = _bounded_days(int(m.group(1) or m.group(2)))
        if result is not None:
            return result

    # ── Tier 3: colloquial timeframes, converted to days ─────────────────────
    if re.search(r"\b(?:immediate\s+joiners?|available\s+immediately|asap|urgently)\b", text_lower):
        return 7   # "immediate" treated as an effective ~1-week ceiling

    m = re.search(r"within\s+(\d{1,2})\s*weeks?\b", text_lower)
    if m:
        result = _bounded_days(int(m.group(1)) * 7)
        if result is not None:
            return result

    m = re.search(r"(\d{1,2})\s*months?\s*(?:'?s)?\s*notice\b", text_lower)
    if m:
        result = _bounded_days(int(m.group(1)) * 30)
        if result is not None:
            return result

    # ── Tier 4: generic catch-all (lowest priority) ──────────────────────────
    for m in re.finditer(cfg.NOTICE_TEXT_EXTRACTION_PATTERN, text_lower):
        window = text_lower[max(0, m.start() - 40): m.end() + 20]
        if "notice" in window or "join" in window or "buy" in window:
            result = _bounded_days(int(m.group(1)))
            if result is not None:
                return result

    # Defensive baseline: nothing matched in any tier. Returning None here
    # is deliberate and safe — every caller treats it as "not stated" and
    # falls back to generic tiers; it is never an unhandled/crashing path.
    return None


def _compute_dynamic_weights(
    category: str, seniority: float, technical_depth: float, urgency: float,
) -> dict:
    """Single-category lookup — kept as a helper for _compute_blended_weights below."""
    base = dict(cfg.CATEGORY_WEIGHT_PROFILES.get(category, cfg.CATEGORY_WEIGHT_PROFILES["default"]))

    sen_delta  = (seniority       - 0.50) * cfg.SENIORITY_ROLE_WEIGHT_BOOST
    tech_delta = (technical_depth - 0.50) * cfg.TECHNICAL_DEPTH_CAP_WEIGHT_BOOST

    base["role_weight"] = float(base["role_weight"]) + sen_delta - tech_delta
    base["cap_weight"]  = float(base["cap_weight"])  - sen_delta + tech_delta

    if urgency > 0.60:
        base["avail_importance"] = float(base["avail_importance"]) + (urgency - 0.60) * cfg.URGENCY_AVAIL_IMPORTANCE_BOOST

    base["role_weight"] = max(0.20, min(0.80, float(base["role_weight"])))
    base["cap_weight"]  = max(0.20, min(0.80, float(base["cap_weight"])))
    total = base["role_weight"] + base["cap_weight"]
    base["role_weight"] /= total
    base["cap_weight"]  /= total

    return base


def _softmax(values: np.ndarray, temperature: float) -> np.ndarray:
    """Standard numerically-stable softmax with temperature scaling."""
    scaled = values / max(temperature, 1e-6)
    scaled = scaled - scaled.max()   # stability
    exp = np.exp(scaled)
    return exp / exp.sum()


def _compute_blended_weights(
    category_sims: dict[str, float],
    seniority: float,
    technical_depth: float,
    urgency: float,
) -> tuple[dict, np.ndarray, dict[str, float]]:
    """
    Replaces winner-take-all category selection. Computes a softmax over
    all category similarities (plus a flat "default" prior), then returns
    the similarity-weighted average of every numeric weight field across
    categories — not a single category's profile.

    Returns (blended_weights_dict, blend_distribution_array_unused_placeholder,
             blend_distribution_dict) for logging/debugging.
    """
    categories = list(cfg.CATEGORY_WEIGHT_PROFILES.keys())   # includes "default"
    sims = np.array(
        [category_sims.get(c, cfg.CATEGORY_BLEND_DEFAULT_PRIOR) for c in categories],
        dtype=np.float64,
    )
    weights = _softmax(sims, cfg.CATEGORY_BLEND_TEMPERATURE)

    # Per-category dynamic weights (after seniority/tech-depth/urgency
    # micro-adjustment), THEN blended — so the micro-adjustments still
    # apply correctly within each category before averaging.
    field_names = ["role_weight", "cap_weight", "avail_importance",
                    "loc_importance", "domain_mismatch_threshold"]
    blended: dict[str, float] = {f: 0.0 for f in field_names}

    for cat, w in zip(categories, weights):
        cat_weights = _compute_dynamic_weights(cat, seniority, technical_depth, urgency)
        for f in field_names:
            blended[f] += w * cat_weights[f]

    # Re-normalise role/cap to sum exactly 1.0 after blending (small drift
    # possible from per-category micro-adjustments before averaging).
    # Epsilon-guarded: role_weight/cap_weight are always positive and
    # bounded per-category (clamped in _compute_dynamic_weights), and the
    # softmax weights always sum to 1.0, so `total` is mathematically
    # guaranteed > 0 — the guard is defensive belt-and-braces, not a fix
    # for a reachable bug.
    total = blended["role_weight"] + blended["cap_weight"]
    if total > 1e-9:
        blended["role_weight"] /= total
        blended["cap_weight"]  /= total

    distribution = {cat: float(w) for cat, w in zip(categories, weights)}
    return blended, weights, distribution


# ─────────────────────────────────────────────────────────────────────────────
# FIX #4 — None-safe formatting helper, used in summary() and logging
# ─────────────────────────────────────────────────────────────────────────────

def _safe_fmt(value: int | float | None, fmt_spec: str = "", na: str = "N/A") -> str:
    """
    Format a possibly-None numeric value safely. Bare f-string interpolation
    of None (f"{x}") happens to print "None" without crashing today, but
    that safety is incidental — it breaks the moment a numeric format spec
    like ":.0f" or ":d" gets added later, which DOES raise on None. Routing
    every potentially-None numeric field through this helper makes the
    safety explicit and permanent rather than accidental, which matters
    under an automated test sandbox running many different JDs in sequence
    where a single crash fails the whole batch.
    """
    if value is None:
        return na
    try:
        return format(value, fmt_spec)
    except (ValueError, TypeError):
        return str(value)


# ─────────────────────────────────────────────────────────────────────────────
# FIX #3 — Domain-mismatch threshold expansion under low confidence
# ─────────────────────────────────────────────────────────────────────────────

# Floor the domain-mismatch threshold can relax to for a genuinely
# out-of-distribution JD. Not zero — the check still provides SOME
# differentiation rather than being fully disabled, it just becomes very
# permissive (cosine sim above ~0.05 against a diffuse blended anchor is a
# trivially low bar for any reasonably professional-sounding candidate).
DOMAIN_MISMATCH_LOW_CONFIDENCE_FLOOR: float = 0.05


def _expand_threshold_for_low_confidence(
    threshold: float, best_sim: float, low_confidence: bool,
) -> tuple[float, float]:
    """
    When the JD's best category match is below CATEGORY_LOW_CONFIDENCE_THRESHOLD,
    relax domain_mismatch_threshold toward DOMAIN_MISMATCH_LOW_CONFIDENCE_FLOOR,
    proportionally to how far below the cutoff best_sim actually sits. A JD
    just barely under the cutoff relaxes only slightly; a JD with near-zero
    (or negative) similarity to every known category relaxes almost fully to
    the floor. Left untouched, an absolute cosine-similarity cutoff checked
    against a near-uninformative blended anchor (the regime "default"
    dominates in) can fail almost every candidate's role vector and
    silently empty the ranked pool for any out-of-distribution JD
    ("Technical Writer", "Chef", "Quantum Cryptographer").

    Returns (possibly-expanded threshold, confidence_ratio used) — the
    ratio is returned purely for logging/transparency.
    """
    if not low_confidence:
        return threshold, 1.0

    confidence_ratio = best_sim / cfg.CATEGORY_LOW_CONFIDENCE_THRESHOLD
    confidence_ratio = max(0.0, min(1.0, confidence_ratio))   # clamps negative best_sim to 0.0 too

    expanded = (
        DOMAIN_MISMATCH_LOW_CONFIDENCE_FLOOR
        + confidence_ratio * (threshold - DOMAIN_MISMATCH_LOW_CONFIDENCE_FLOOR)
    )
    return expanded, confidence_ratio


# ─────────────────────────────────────────────────────────────────────────────
# JDProcessor
# ─────────────────────────────────────────────────────────────────────────────

class JDProcessor:
    """
    Processes any JD (.docx/.txt/.md) and produces JDProfile + 4 formats.
    Category weighting is a blend across all known categories (see module
    docstring) — no JD-specific hardcoding, no winner-take-all bucket cutoff.
    """

    def __init__(self, jd_path: str | Path) -> None:
        self._path    = Path(jd_path)
        self._raw:    str | None        = None
        self._flags:  dict | None       = None
        self._rvec:   np.ndarray | None = None
        self._cvec:   np.ndarray | None = None
        self._rkw:    list[str] | None  = None
        self._ckw:    list[str] | None  = None
        self._profile:JDProfile | None  = None
        self._model   = None
        _ = self.raw_text

    # ── Format D ──────────────────────────────────────────────────────────────
    @property
    def raw_text(self) -> str:
        if self._raw is None:
            self._raw = self._ingest()
            log.info("JD ingested: %d chars from '%s'", len(self._raw), self._path.name)
        return self._raw

    def _ingest(self) -> str:
        suffix = self._path.suffix.lower()
        if suffix == ".docx":
            return self._read_docx()
        try:
            return self._path.read_text(encoding="utf-8")
        except Exception as e:
            raise ValueError(f"Cannot read '{self._path}': {e}") from e

    def _read_docx(self) -> str:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("pip install python-docx")
        doc   = Document(self._path)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return "\n".join(lines)

    # ── Format A ──────────────────────────────────────────────────────────────
    @property
    def flags(self) -> dict[str, bool]:
        if self._flags is None:
            self._flags = cfg.extract_jd_flags(self.raw_text)
            log.info("JD flags: %d true", sum(self._flags.values()))
        return self._flags

    # ── Embedding ─────────────────────────────────────────────────────────────
    def _load_model(self):
        if self._model is None:
            from model_engine import ONNXEmbedder
            path = str(cfg.BGE_EMBED_MODEL_DIR) if cfg.BGE_EMBED_MODEL_DIR.exists() else cfg.BGE_EMBED_MODEL_ID
            log.info("Loading embed model for JD encoding …")
            self._model = ONNXEmbedder(cfg.BGE_EMBED_MODEL_DIR, fallback_model_id=cfg.BGE_EMBED_MODEL_ID)
            log.info("JD embed model backend: %s", "ONNX INT8" if self._model.using_onnx else "PyTorch (fallback)")
        return self._model

    def _encode(self, text: str) -> np.ndarray:
        return self._load_model().encode(text, normalize_embeddings=True, convert_to_numpy=True).astype(np.float32)

    # ── Section split ─────────────────────────────────────────────────────────
    def _split_sections(self) -> tuple[str, str]:
        lines = self.raw_text.split("\n")
        role_l: list[str] = []
        cap_l:  list[str] = []
        cur: str | None   = None
        for line in lines:
            ll = line.lower().strip()
            if any(h in ll for h in cfg.ROLE_SECTION_HEADERS):       cur = "role"
            elif any(h in ll for h in cfg.CAPABILITY_SECTION_HEADERS): cur = "cap"
            if cur == "role": role_l.append(line)
            elif cur == "cap": cap_l.append(line)
        return ("\n".join(role_l).strip() or self.raw_text[:500],
                "\n".join(cap_l).strip()  or self.raw_text)

    # ── Format B ──────────────────────────────────────────────────────────────
    @property
    def role_vector(self) -> np.ndarray:
        if self._rvec is None: self._compute_vecs()
        return self._rvec   # type: ignore

    @property
    def cap_vector(self) -> np.ndarray:
        if self._cvec is None: self._compute_vecs()
        return self._cvec   # type: ignore

    def _compute_vecs(self) -> None:
        r, c = self._split_sections()
        log.info("Encoding JD: role=%d cap=%d chars", len(r), len(c))
        self._rvec = self._encode(r)
        self._cvec = self._encode(c)

    # ── Format C ──────────────────────────────────────────────────────────────
    @property
    def role_keywords(self) -> list[str]:
        if self._rkw is None: self._rkw = self._tok(self._split_sections()[0])
        return self._rkw

    @property
    def cap_keywords(self) -> list[str]:
        if self._ckw is None: self._ckw = self._tok(self._split_sections()[1])
        return self._ckw

    @property
    def keywords(self) -> list[str]:
        return self._tok(self.raw_text)

    @staticmethod
    def _tok(text: str) -> list[str]:
        text = re.sub(r"[^a-z0-9\s\-]", " ", text.lower())
        seen: set[str] = set(); result: list[str] = []
        for t in text.split():
            t = t.strip("-")
            if len(t) > 2 and t not in cfg.STOPWORDS and t not in seen:
                seen.add(t); result.append(t)
        return result

    # ── JDProfile (core of scalability) ──────────────────────────────────────
    @property
    def profile(self) -> JDProfile:
        if self._profile is None:
            self._profile = self._build_profile()
        return self._profile

    def _build_profile(self) -> JDProfile:
        text_lower = self.raw_text.lower()

        seniority       = _detect_seniority(text_lower)
        technical_depth = _detect_technical_depth(text_lower)
        urgency         = _detect_urgency(text_lower)
        log.info("JD detect: seniority=%.2f tech=%.2f urgency=%.2f", seniority, technical_depth, urgency)

        # ── Category similarities — computed against ALL anchors, no argmax cut ──
        jd_role_vec = self.role_vector
        category_sims, anchor_vecs = self._compute_category_similarities(jd_role_vec)

        best_cat = max(category_sims, key=category_sims.get)
        best_sim = category_sims[best_cat]
        low_confidence = best_sim < cfg.CATEGORY_LOW_CONFIDENCE_THRESHOLD
        log.info(
            "JD category (highest-sim, display only): '%s' (sim=%.3f)%s",
            best_cat, best_sim,
            "  [LOW CONFIDENCE — no category resembles this JD well, blend regresses toward default]" if low_confidence else "",
        )

        # ── Blend weights across ALL categories (this is what's actually used) ──
        blended, blend_weights, distribution = _compute_blended_weights(
            category_sims, seniority, technical_depth, urgency,
        )
        top3 = sorted(distribution.items(), key=lambda kv: -kv[1])[:3]
        log.info("JD category blend (top 3 of %d): %s", len(distribution),
                 ", ".join(f"{c}={w:.2f}" for c, w in top3))

        # ── FIX #3: expand domain_mismatch_threshold under low confidence ──────
        expanded_threshold, confidence_ratio = _expand_threshold_for_low_confidence(
            blended["domain_mismatch_threshold"], best_sim, low_confidence,
        )
        if low_confidence:
            log.info(
                "Low-confidence domain match: expanding domain_mismatch_threshold "
                "%.3f -> %.3f (confidence_ratio=%.2f) to avoid over-filtering an "
                "out-of-distribution JD",
                blended["domain_mismatch_threshold"], expanded_threshold, confidence_ratio,
            )
        blended["domain_mismatch_threshold"] = expanded_threshold

        # Blended anchor vector — weighted average of all anchor embeddings,
        # same softmax weights, used for the domain-mismatch check
        blended_anchor = np.zeros(cfg.EMBED_DIM, dtype=np.float32)
        for (cat, _sim), w in zip(category_sims.items(), blend_weights):
            blended_anchor += w * anchor_vecs[cat]
        norm = np.linalg.norm(blended_anchor)
        if norm > 1e-9:
            blended_anchor = (blended_anchor / norm).astype(np.float32)

        dq = _extract_dq_conditions(text_lower)
        preferred_cities = _extract_preferred_cities(self.raw_text)
        log.info("JD preferred cities (auto): %s", sorted(preferred_cities))

        # ── FIX #1: widened multi-range YoE extraction ──────────────────────────
        yoe_min, yoe_max = _detect_yoe_range(self.raw_text)
        if yoe_min is not None or yoe_max is not None:
            log.info("JD YoE range: %s years", _format_yoe_range(yoe_min, yoe_max))

        nlp_ir_required = _detect_nlp_ir_required(text_lower)

        # ── FIX #2: notice-period cascade ───────────────────────────────────────
        notice_stated = _extract_notice_period_days(self.raw_text)
        if notice_stated is not None:
            log.info("JD states explicit notice period: %s days (used instead of fallback tiers)",
                      _safe_fmt(notice_stated, "d"))

        return JDProfile(
            role_category=best_cat,
            category_confidence=best_sim,
            category_low_confidence=low_confidence,
            seniority_level=seniority, technical_depth=technical_depth, urgency=urgency,
            role_weight=float(blended["role_weight"]), cap_weight=float(blended["cap_weight"]),
            avail_importance=float(blended["avail_importance"]),
            loc_importance=float(blended["loc_importance"]),
            domain_mismatch_threshold=float(blended["domain_mismatch_threshold"]),
            # text-native, no category-default fallback (see module docstring)
            consulting_penalty_active=dq["consulting_penalty_active"],
            research_penalty_active=dq["research_penalty_active"],
            production_is_required=dq["production_is_required"],
            preferred_cities=preferred_cities,
            yoe_min=yoe_min, yoe_max=yoe_max,
            nlp_ir_required=nlp_ir_required,
            notice_period_days_stated=notice_stated,
            category_anchor_vec=blended_anchor,
            flags=self.flags,
        )

    def _compute_category_similarities(
        self, jd_role_vec: np.ndarray,
    ) -> tuple[dict[str, float], dict[str, np.ndarray]]:
        """
        Cosine similarity of the JD's role vector against every category
        anchor (the real 10, not "default" — "default" gets a flat prior
        instead of an embedding, see _compute_blended_weights)
        """
        model = self._load_model()
        sims: dict[str, float] = {}
        anchor_vecs: dict[str, np.ndarray] = {}
        for cat_name, anchor_text in cfg.ROLE_CATEGORY_ANCHORS.items():
            av = model.encode(anchor_text, normalize_embeddings=True, convert_to_numpy=True).astype(np.float32)
            sims[cat_name] = float(np.dot(jd_role_vec, av))
            anchor_vecs[cat_name] = av
        # "default" has no embedding of its own — give it a neutral
        # zero-vector so it contributes weight to the blend (per its prior)
        # without skewing the anchor DIRECTION average toward any specific
        # semantic content. The flat prior in _compute_blended_weights
        # determines how much pull it gets, not a fake embedding here.
        anchor_vecs["default"] = np.zeros(cfg.EMBED_DIM, dtype=np.float32)
        sims["default"] = cfg.CATEGORY_BLEND_DEFAULT_PRIOR
        return sims, anchor_vecs

    def summary(self) -> dict[str, Any]:
        """
        Human-readable summary. FIX #4: every potentially-None numeric
        field is routed through _safe_fmt rather than bare f-string
        interpolation, so this can never crash regardless of which JD
        (including edge cases that legitimately produce None for yoe_min
        or notice_period_days_stated) is fed through an automated test loop.
        """
        p = self.profile
        return {
            "jd_file": self._path.name,
            "role_category": p.role_category,
            "confidence": _safe_fmt(p.category_confidence, ".3f"),
            "low_confidence_blend": p.category_low_confidence,
            "seniority": _safe_fmt(p.seniority_level, ".2f"),
            "tech_depth": _safe_fmt(p.technical_depth, ".2f"),
            "urgency": _safe_fmt(p.urgency, ".2f"),
            "weights": {
                "role": _safe_fmt(p.role_weight, ".2f"),
                "cap": _safe_fmt(p.cap_weight, ".2f"),
            },
            "domain_mismatch_threshold": _safe_fmt(p.domain_mismatch_threshold, ".3f"),
            "preferred_cities": sorted(p.preferred_cities),
            "yoe_range": _format_yoe_range(p.yoe_min, p.yoe_max),
            "nlp_ir_required": p.nlp_ir_required,
            "notice_period_days_stated": _safe_fmt(p.notice_period_days_stated, "d"),
            "penalty_active": {"consulting": p.consulting_penalty_active, "research": p.research_penalty_active},
            "flags_true": [k for k, v in p.flags.items() if v],
        }